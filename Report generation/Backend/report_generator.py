"""PDF and DOCX report generation for student performance."""

from pathlib import Path

from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
from docx import Document


def generate_student_report(student, reports_dir):
    """Create a professional PDF report and return its filename."""
    reports_path = Path(reports_dir)
    reports_path.mkdir(parents=True, exist_ok=True)
    filename = f"student_report_{student['student_id']}.pdf"
    file_path = reports_path / filename

    document = SimpleDocTemplate(
        str(file_path),
        pagesize=A4,
        rightMargin=0.65 * inch,
        leftMargin=0.65 * inch,
        topMargin=0.65 * inch,
        bottomMargin=0.65 * inch,
    )
    styles = getSampleStyleSheet()
    title_style = styles["Title"]
    title_style.alignment = TA_CENTER

    details = [
        ["Student Name", student["student_name"]],
        ["Student ID", str(student["student_id"])],
        ["Department", student["department"]],
        ["Company", student["company"]],
        ["Attendance", f"{student['attendance']}%"],
        ["Task Completion", f"{student['task_completion']}%"],
        ["Mentor Rating", f"{student['mentor_rating']} / 5"],
        ["Communication", f"{student['communication']}%"],
        ["Project Completion", f"{student['project_completion']}%"],
        ["Overall Score", f"{student['overall_score']}%"],
        ["Performance Category", student["performance_category"]],
    ]
    table = Table(details, colWidths=[2.1 * inch, 4.25 * inch])
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#EEEAFE")),
                ("TEXTCOLOR", (0, 0), (0, -1), colors.HexColor("#4C1D95")),
                ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#D1D5DB")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("PADDING", (0, 0), (-1, -1), 9),
            ]
        )
    )
    document.build(
        [
            Paragraph("AI Internship Management System", title_style),
            Spacer(1, 0.15 * inch),
            Paragraph("Student Performance Report", styles["Heading2"]),
            Spacer(1, 0.2 * inch),
            table,
        ]
    )
    return filename


def report_summary(students):
    """Return aggregate values shared by the report preview and downloads."""
    scores = [float(student["overall_score"]) for student in students]
    return {
        "average_score": round(sum(scores) / len(scores), 2) if scores else 0,
        "highest_score": round(max(scores), 2) if scores else 0,
        "lowest_score": round(min(scores), 2) if scores else 0,
        "outstanding_students": sum(s["performance_category"] == "Outstanding" for s in students),
        "needs_improvement_students": sum(s["performance_category"] == "Needs Improvement" for s in students),
    }


def _report_rows(students):
    return [[
        student["student_name"], str(student["student_id"]), student["department"], student["company"],
        f"{student['attendance']}%", f"{student['task_completion']}%", f"{student['communication']}%",
        f"{student['mentor_rating']}/5", f"{student['project_completion']}%", f"{student['overall_score']}%",
        student["performance_category"],
    ] for student in students]


def generate_range_pdf(students, start_date, end_date, reports_dir):
    """Generate a dated PDF report for all students in the selected range."""
    reports_path = Path(reports_dir); reports_path.mkdir(parents=True, exist_ok=True)
    filename = f"performance_report_{start_date}_to_{end_date}.pdf"; path = reports_path / filename
    document = SimpleDocTemplate(str(path), pagesize=A4, rightMargin=22, leftMargin=22, topMargin=28, bottomMargin=28)
    styles = getSampleStyleSheet(); title = styles["Title"]; title.alignment = TA_CENTER
    summary = report_summary(students)
    headers = ["Student", "ID", "Dept.", "Company", "Attend.", "Tasks", "Comm.", "Mentor", "Project", "Score", "Category"]
    table = Table([headers, *_report_rows(students)] if students else [headers, ["No records in selected date range", *[""] * 10]], repeatRows=1, colWidths=[56, 28, 46, 50, 38, 34, 38, 34, 38, 34, 56])
    table.setStyle(TableStyle([("BACKGROUND", (0,0), (-1,0), colors.HexColor("#8B5CF6")), ("TEXTCOLOR", (0,0), (-1,0), colors.white), ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"), ("FONTSIZE", (0,0), (-1,-1), 6.5), ("GRID", (0,0), (-1,-1), .25, colors.HexColor("#CBD5E1")), ("VALIGN", (0,0), (-1,-1), "MIDDLE"), ("PADDING", (0,0), (-1,-1), 4), ("BACKGROUND", (0,1), (-1,-1), colors.white)]))
    document.build([Paragraph("AI Internship Management System", title), Spacer(1, 8), Paragraph("Student Performance Report", styles["Heading2"]), Paragraph(f"Selected date range: {start_date} to {end_date}", styles["Normal"]), Paragraph(f"Average score: {summary['average_score']}%  |  Highest: {summary['highest_score']}%  |  Lowest: {summary['lowest_score']}%  |  Outstanding: {summary['outstanding_students']}  |  Needs improvement: {summary['needs_improvement_students']}", styles["Normal"]), Spacer(1, 10), table])
    return filename


def generate_range_docx(students, start_date, end_date, reports_dir):
    """Generate an editable DOCX report for the selected date range."""
    reports_path = Path(reports_dir); reports_path.mkdir(parents=True, exist_ok=True)
    filename = f"performance_report_{start_date}_to_{end_date}.docx"; path = reports_path / filename
    document = Document(); document.add_heading("AI Internship Management System", 0); document.add_heading("Student Performance Report", 1)
    document.add_paragraph(f"Selected date range: {start_date} to {end_date}")
    summary = report_summary(students); document.add_paragraph(f"Average Score: {summary['average_score']}% | Highest Score: {summary['highest_score']}% | Lowest Score: {summary['lowest_score']}% | Outstanding Students: {summary['outstanding_students']} | Needs Improvement Students: {summary['needs_improvement_students']}")
    headers = ["Student", "ID", "Department", "Company", "Attendance", "Tasks", "Communication", "Mentor", "Project", "Score", "Category"]
    table = document.add_table(rows=1, cols=len(headers)); table.style = "Table Grid"
    for index, header in enumerate(headers): table.rows[0].cells[index].text = header
    for row in _report_rows(students):
        cells = table.add_row().cells
        for index, value in enumerate(row): cells[index].text = value
    if not students: document.add_paragraph("No records in selected date range.")
    document.save(path); return filename
